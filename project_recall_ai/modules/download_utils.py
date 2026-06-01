import pandas as pd
import json
from io import BytesIO
from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfgen import canvas
from docx import Document


# ===============================
# CSV
# ===============================
def export_csv(df: pd.DataFrame, summary: str) -> bytes:
    buffer = BytesIO()
    buffer.write(f"# SUMMARY\n# {summary.replace('\n', ' ')}\n\n".encode())
    df.to_csv(buffer, index=False)
    return buffer.getvalue()


def export_json(df: pd.DataFrame, summary: str) -> bytes:
    payload = {
        "summary": summary,
        "records": df.fillna("").to_dict(orient="records"),
    }
    return json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")


# ===============================
# EXCEL
# ===============================
def export_excel(df: pd.DataFrame, summary: str) -> bytes:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="Results", index=False)

        summary_df = pd.DataFrame({"Summary": summary.split("\n")})
        summary_df.to_excel(writer, sheet_name="Summary", index=False)

    return buffer.getvalue()


# ===============================
# PDF
# ===============================
def export_pdf(df: pd.DataFrame, summary: str) -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    width, height = landscape(A4)

    y = height - 40
    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, y, "AI Query Summary")
    y -= 30

    c.setFont("Helvetica", 10)
    for line in summary.split("\n"):
        c.drawString(40, y, line)
        y -= 14
        if y < 60:
            c.showPage()
            y = height - 40

    if df.empty:
        c.save()
        return buffer.getvalue()

    c.showPage()
    y = height - 40
    c.setFont("Helvetica-Bold", 12)
    c.drawString(40, y, "Results")
    y -= 20
    c.setFont("Helvetica", 8)

    display_cols = list(df.columns[:8])
    if len(df.columns) > len(display_cols):
        c.drawString(40, y, f"Showing first {len(display_cols)} of {len(df.columns)} columns for PDF readability.")
        y -= 14

    header = " | ".join(display_cols)
    c.drawString(40, y, header[:150])
    y -= 14

    for _, row in df.head(40).iterrows():
        row_text = " | ".join(str(row.get(col, ""))[:28] for col in display_cols)
        c.drawString(40, y, row_text)
        y -= 12
        if y < 40:
            c.showPage()
            y = height - 40

    c.save()
    return buffer.getvalue()


# ===============================
# WORD
# ===============================
def export_word(df: pd.DataFrame, summary: str) -> bytes:
    doc = Document()
    doc.add_heading("AI Query Summary", level=1)

    if summary.strip():
        for line in summary.split("\n"):
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("No AI summary was generated for this report.")

    if df.empty:
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    doc.add_heading("Results Table", level=2)

    if len(df.columns) > 8:
        doc.add_paragraph(
            f"The result contains {len(df.columns)} columns. Records are shown as field lists for readability."
        )
        for idx, row in df.head(40).iterrows():
            doc.add_heading(f"Record {idx + 1}", level=3)
            for col, val in row.items():
                doc.add_paragraph(f"{col}: {val}")
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    table = doc.add_table(rows=1, cols=max(len(df.columns), 1))
    table.style = "Table Grid"
    hdrs = table.rows[0].cells

    for i, col in enumerate(df.columns):
        hdrs[i].text = col

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
